from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader

from app.config import get_settings
from app.schemas import FileExtractResponse, SourceType
from app.services.source_loader import clean_text


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


def _source_type_from_extension(extension: str) -> SourceType:
    if extension == ".pdf":
        return "pdf"
    if extension == ".docx":
        return "docx"
    return "text"


async def extract_upload_text(file: UploadFile) -> FileExtractResponse:
    settings = get_settings()
    filename = file.filename or "upload"
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Supported file types: .txt, .pdf, .docx. Image OCR is planned for a later version.",
        )

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="File is larger than 10MB.")

    warnings: list[str] = []
    if extension == ".txt":
        text = content.decode("utf-8", errors="ignore")
    elif extension == ".pdf":
        reader = PdfReader(BytesIO(content))
        pages = [(page.extract_text() or "") for page in reader.pages]
        text = "\n".join(pages)
        if not clean_text(text):
            warnings.append("No selectable text found in PDF. It may be scanned and require OCR.")
    else:
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells: list[str] = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_cells.append(cell.text)
        text = "\n".join([*paragraphs, *table_cells])

    cleaned = clean_text(text, max_chars=settings.max_input_chars)
    if len(cleaned) >= settings.max_input_chars:
        warnings.append(f"Extracted text was truncated to {settings.max_input_chars} characters.")

    return FileExtractResponse(
        filename=filename,
        source_type=_source_type_from_extension(extension),
        text=cleaned,
        character_count=len(cleaned),
        warnings=warnings,
    )

